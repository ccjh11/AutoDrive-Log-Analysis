import os
import numpy as np
import matplotlib.pyplot as plt
from core.data_processor import load_and_clean_log
from core.rule_engine import RuleEngine, generate_report

from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

def plot_signals_and_save(frames, output_dir="output_figures"):
    """
    生成曲线: 车速(Speed)、目标距离(Target_Distance)、AEB触发(AEB_Trigger)、制动压力(Brake_Pressure)、DTC码(DTC_Code)
    返回每个曲线的图片路径dict
    画图部分每10个点取1个点采样绘图以提升速度
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 支持frame是dict或对象
    def extract_value(f, key):
        # 优先用key，不区分大小写
        if isinstance(f, dict):
            # 支持不区分大小写字段
            for k in f:
                if k.lower() == key.lower():
                    return f[k]
            return None
        else:
            # 尝试直接getattr，或小写
            if hasattr(f, key):
                return getattr(f, key)
            elif hasattr(f, key.lower()):
                return getattr(f, key.lower())
            else:
                # 还可以尝试属性名全小写
                for attr in dir(f):
                    if attr.lower() == key.lower():
                        return getattr(f, attr)
                return None

    # 采用以上extract_value，确保字段妥善读取
    times = []
    speeds = []
    distances = []
    aeb_triggers = []
    brake_pressures = []
    dtc_codes = []

    for f in frames:
        times.append(extract_value(f, 'time'))
        speeds.append(extract_value(f, 'Speed'))
        distances.append(extract_value(f, 'Target_Distance'))
        aeb = extract_value(f, 'AEB_Trigger')
        # 如果AEB_Trigger为None自动转为0
        aeb_triggers.append(0 if aeb is None else aeb)
        bp = extract_value(f, 'Brake_Pressure')
        brake_pressures.append(0 if bp is None else bp)
        dtc_codes.append(extract_value(f, 'DTC_Code'))

    # 转np.array，过滤掉None（只对坐标轴进行简单处理）
    times_array = np.array([x for x in times if x is not None])
    speeds_array = np.array([x if x is not None else np.nan for x in speeds])
    distances_array = np.array([x if x is not None else np.nan for x in distances])
    aeb_triggers_array = np.array([x if x is not None else 0 for x in aeb_triggers])
    brake_pressures_array = np.array([x if x is not None else np.nan for x in brake_pressures])
    # dtc_codes 保持原始顺序

    # 如果所有数据都为空，给一个空数组，避免报错
    if len(times_array) == 0:
        times_array = np.arange(len(frames))
    paths = {}

    # ====== 加采样：只取每10个数据 ======
    step = 10
    times_sample = times_array[::step]
    speeds_sample = speeds_array[::step]
    distances_sample = distances_array[::step]
    aeb_triggers_sample = aeb_triggers_array[::step]
    brake_pressures_sample = brake_pressures_array[::step]
    # 注：DTC码文本原样导出，不采样

    # 车速曲线
    plt.figure(figsize=(10, 4))
    plt.plot(times_sample, speeds_sample, label="Speed (km/h)", color='b')
    plt.xlabel("Time (s)")
    plt.ylabel("Speed (km/h)")
    plt.title("Vehicle Speed Curve")
    plt.grid(True)
    plt.tight_layout()
    speed_path = os.path.join(output_dir, "speed_curve.png")
    plt.savefig(speed_path)
    plt.close()
    paths['Speed'] = speed_path

    # 目标距离曲线
    plt.figure(figsize=(10, 4))
    plt.plot(times_sample, distances_sample, label="Target Distance (m)", color='g')
    plt.xlabel("Time (s)")
    plt.ylabel("Target Distance (m)")
    plt.title("Target Distance Curve")
    plt.grid(True)
    plt.tight_layout()
    distance_path = os.path.join(output_dir, "target_distance_curve.png")
    plt.savefig(distance_path)
    plt.close()
    paths['Target_Distance'] = distance_path

    # AEB触发状态曲线
    plt.figure(figsize=(10, 2))
    plt.step(times_sample, aeb_triggers_sample, label="AEB Trigger", color='r', where="post")
    plt.xlabel("Time (s)")
    plt.ylabel("AEB Trigger")
    plt.title("AEB Trigger Status Curve")
    plt.yticks([0, 1], ["Off", "On"])
    plt.grid(True, which='both', axis='x')
    plt.tight_layout()
    aeb_path = os.path.join(output_dir, "aeb_trigger_curve.png")
    plt.savefig(aeb_path)
    plt.close()
    paths['AEB_Trigger'] = aeb_path

    # 制动压力曲线
    plt.figure(figsize=(10, 4))
    plt.plot(times_sample, brake_pressures_sample, label="Brake Pressure", color='m')
    plt.xlabel("Time (s)")
    plt.ylabel("Brake Pressure")
    plt.title("Brake Pressure Curve")
    plt.grid(True)
    plt.tight_layout()
    brake_pressure_path = os.path.join(output_dir, "brake_pressure_curve.png")
    plt.savefig(brake_pressure_path)
    plt.close()
    paths['Brake_Pressure'] = brake_pressure_path

    # DTC Code 序列文本输出, 用表格而不是曲线呈现
    dtc_code_txt_path = os.path.join(output_dir, "dtc_code_list.txt")
    with open(dtc_code_txt_path, "w", encoding="utf-8") as f:
        for t, dtc in zip(times, dtc_codes):
            f.write(f"{t},{dtc}\n")
    paths['DTC_Code'] = dtc_code_txt_path

    return paths

def html_escape(text):
    """简单过滤HTML中的特殊字符"""
    return (str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;"))

def generate_html_report(
        report_title,
        run_time,
        rule_reports,
        signal_figs,
        exception_figs):
    """
    自动生成html形式的报告
    - report_title: 测试名称
    - run_time: 测试时间
    - rule_reports: 自动化检测结果（list of dict）
    - signal_figs: dict of curve png paths
    - exception_figs: (可选)每个异常的截图，list of png path
    """

    html_lines = []
    html_lines.append(f'<h1>{html_escape(report_title)}</h1>')
    html_lines.append(f'<p><b>测试时间：</b>{html_escape(run_time)}<br>')
    html_lines.append(f'<b>异常总数：</b>{len(rule_reports)}<br></p>')

    # 曲线插图
    html_lines.append('<h2>测试主要信号曲线</h2>')
    # 参数名与展示
    param_label_map = {
        'Speed': "车速(Speed)",
        'Target_Distance': "目标距离(Target_Distance)",
        'AEB_Trigger': "AEB触发(AEB_Trigger)",
        'Brake_Pressure': "制动压力(Brake_Pressure)",
        'DTC_Code': "DTC故障码(DTC_Code)"
    }
    for label, path in signal_figs.items():
        display_label = html_escape(param_label_map.get(label, label))
        if label == 'DTC_Code':
            html_lines.append(f'<div><b>{display_label}:</b></div>')
            html_lines.append(f'<a href="{path}" target="_blank">下载DTC故障码表</a><br>')
        else:
            html_lines.append(f'<div><b>{display_label}:</b></div>')
            html_lines.append(f'<img src="{path}" width="720"/><br>')

    # 异常详情
    html_lines.append('<h2>异常检测与分析</h2>')
    if not rule_reports:
        html_lines.append("<p>所有检查通过，未发现异常。</p>")
    else:
        for idx, rep in enumerate(rule_reports, 1):
            risk = html_escape(rep.get("risk", ""))
            fault = html_escape(rep.get("type", ""))
            desc = html_escape(rep.get("desc", ""))
            cause = html_escape(rep.get("cause", ""))
            time = html_escape(rep.get("time", ""))
            html_lines.append(f'<div style="border:1px solid #ddd; margin:10px 0; padding:8px;">')
            html_lines.append(f'<b>{idx}. [{risk}] {fault} @ {time}</b><br>')
            html_lines.append(f'<span>说明：{desc}</span><br>')
            if cause:
                html_lines.append(f'原因：{cause}<br>')
            # 异常截图（如有）
            if exception_figs and idx-1 < len(exception_figs):
                html_lines.append(f'<img src="{exception_figs[idx-1]}" width="400"/><br>')
            html_lines.append('</div>')

    return '\n'.join(html_lines)

def main(
    log_path,
    report_title="台架自动化日志分析报告",
    output_dir="output_result"
):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    print(f"[系统] 正在加载日志：{log_path}")
    frames = load_and_clean_log(log_path)
    print(f"[系统] 日志帧加载完成，共 {len(frames)} 条")
    # 1. 画图 & 保存
    figs = plot_signals_and_save(frames, output_dir=os.path.join(output_dir, "figures"))

    # 2. 执行规则检测
    engine = RuleEngine(frames)
    rules_report = engine.analyze()
    print(f"[系统] 自动检测完成，异常数量: {len(rules_report)}")

    # 3. 每个异常可进一步做局部截图/可选，本例仅输出全图
    exception_figs = []  # 可扩展，每个异常生成专属截图

    # 4. 输出html报告
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    html = generate_html_report(
        report_title=report_title,
        run_time=now_str,
        rule_reports=rules_report,
        signal_figs=figs,
        exception_figs=exception_figs
    )
    html_path = os.path.join(output_dir, "report.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"[系统] HTML测试报告已生成：{html_path}")

    # 5. 生成Word报告（含主要曲线和异常摘要）
    # 限制 Word 报告遍历异常只输出前100条
    try:
        doc = Document()
        doc.add_heading(report_title, 0)
        doc.add_paragraph(f"测试时间：{now_str}")
        doc.add_paragraph(f"异常总数：{len(rules_report)}")

        doc.add_heading("测试主要信号曲线", level=1)
        param_label_map = {
            'Speed': "车速(Speed)",
            'Target_Distance': "目标距离(Target_Distance)",
            'AEB_Trigger': "AEB触发(AEB_Trigger)",
            'Brake_Pressure': "制动压力(Brake_Pressure)",
            'DTC_Code': "DTC故障码(DTC_Code)"
        }
        for label, path in figs.items():
            display_label = param_label_map.get(label, label)
            if label == 'DTC_Code':
                doc.add_paragraph(f"{display_label}：见文件 {os.path.basename(path)}")
            else:
                doc.add_paragraph(f"{display_label}曲线：")
                doc.add_picture(path, width=Inches(6.0))

        doc.add_heading("异常检测与分析", level=1)
        if not rules_report:
            doc.add_paragraph("所有检查通过，未发现异常。")
        else:
            # 只输出前100条异常
            for idx, rep in enumerate(rules_report[:100], 1):
                doc.add_heading(f"{idx}. [{rep.get('risk', '')}] {rep.get('type', '')} @ {rep.get('time', '')}", level=2)
                doc.add_paragraph(rep.get('desc', ''))
                cause = rep.get('cause', '')
                if cause:
                    doc.add_paragraph(f"原因：{cause}")
                evidence = rep.get('evidence', '')
                if evidence:
                    doc.add_paragraph(f"证据关键：{str(evidence)}")
            if len(rules_report) > 100:
                doc.add_paragraph(f"...... 共{len(rules_report)}条，仅显示前100条")
        word_path = os.path.join(output_dir, "report.docx")
        doc.save(word_path)
        print(f"[系统] Word测试报告已生成：{word_path}")
    except Exception as e:
        print(f"[警告] 生成Word报告失败: {e}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="自动日志分析系统")
    parser.add_argument("--log", help="日志csv文件路径", required=True)
    parser.add_argument("--name", help="测试名称", default="台架自动化日志分析报告")
    parser.add_argument("--out", help="结果输出目录", default="output_result")
    args = parser.parse_args()
    main(log_path=args.log, report_title=args.name, output_dir=args.out)